"""
완성본 — docs/제출본_최종_임과함께.docx (+ .pdf)

붙임3 데이터 분석 결과서 본문에 소스코드 전문을 별첨으로 이어 붙입니다.
요강이 "코드 소스 등 분석 결과물 필수 첨부(분량 외 별첨)"라고 하므로 별첨은
쪽수에 들어가지 않습니다. 본문은 그대로 10쪽입니다.

만드는 방식은 make_docx.py가 만든 문서를 열어 뒤에 덧붙이는 것입니다. 본문을
여기서 다시 짜면 두 벌이 생기고, 한쪽만 고치는 일이 반드시 따라옵니다.

PDF는 만들어진 docx의 XML을 다시 읽어 HTML로 옮긴 뒤 Chromium으로 인쇄합니다.
이 컨테이너에 워드도 리브레오피스도 없어 다른 길이 없기도 하지만, 그 덕에
"내가 의도한 것"이 아니라 "파일에 실제로 들어간 것"을 인쇄하게 됩니다.

실행: python src/make_final_docx.py [--no-pdf]
"""
from __future__ import annotations

import argparse
import os
import subprocess
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import make_docx as M

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
BASE = os.path.join(ROOT, "docs", "제출본_임과함께_데이터분석부문.docx")
OUT = os.path.join(ROOT, "docs", "제출본_최종_임과함께.docx")
OUT_PDF = os.path.join(ROOT, "docs", "제출본_최종_임과함께.pdf")

CODE_FONT = "Consolas"
CODE_INK = RGBColor(0x24, 0x2C, 0x31)

# 별첨에 실을 코드. 방법론을 담은 것부터 놓아, 앞에서부터 읽어도 흐름이 잡히게 합니다.
GROUPS = [
    ("가. 전처리 — 자료를 다듬고 누수를 막는다", [
        "src/preprocess.py", "src/preprocess_cost.py"]),
    ("나. 학습 — 하이퍼파라미터 탐색과 벤치마크", [
        "src/train_optuna.py", "src/train_cost.py",
        "src/train_quantile.py", "src/train_panel.py", "src/tune_panel.py"]),
    ("다. 검증 — 스스로 의심하고 확인한다", [
        "src/audit_split.py"]),
    ("라. 설명과 처방", [
        "src/explain.py", "src/portfolio.py"]),
    ("마. 분석 계층", [
        "src/insights.py", "src/production.py", "src/management.py",
        "src/subsidy.py", "src/region_map.py"]),
    ("바. 외부 자료 결합", [
        "src/kamis_client.py", "src/shipping.py",
        "src/kma_client.py", "src/weather.py", "src/weather_sgg.py"]),
    ("사. 서비스", [
        "api/main.py", "api/services.py"]),
]


def code_para(doc, text, size=7.4, color=CODE_INK, shade=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.06
    pf.left_indent = Cm(0.25)
    M._font(p.add_run(text if text.strip() else " "), size, False, color,
            CODE_FONT, M.FONT)
    if shade:
        M._shade(p, shade)
    return p


def add_code_file(doc, rel: str) -> int:
    path = os.path.join(ROOT, rel)
    if not os.path.exists(path):
        print(f"  [건너뜀] {rel}")
        return 0
    lines = open(path, encoding="utf-8").read().rstrip("\n").split("\n")

    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(11)
    h.paragraph_format.space_after = Pt(3)
    h.paragraph_format.keep_with_next = True
    M._font(h.add_run(f"▸ {rel}"), 9.5, True, M.FOREST)
    M._font(h.add_run(f"   ({len(lines):,}줄)"), 8, False, M.MUTED)

    for ln in lines:
        # 탭은 문서에서 폭이 들쭉날쭉해집니다. 공백 넷으로 폅니다.
        code_para(doc, ln.replace("\t", "    "), shade="F6F8F9")
    return len(lines)


def append_appendix(doc) -> tuple[int, int]:
    doc.add_page_break()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    M._font(p.add_run("별첨. 분석 코드 전문"), 15, True, M.FOREST)
    M._border(p, "bottom", "C8D8CE", 8)

    M.P(doc, "요강 붙임3 유의사항 “프로그래밍 등 이용 분석 사례의 경우, 코드 소스 등 "
             "분석 결과물 필수 첨부(분량 외 별첨)”에 따른 별첨입니다. **본문 분량에는 "
             "들어가지 않습니다.**")
    M.NOTE(doc, "**원자료는 함께 싣지 않았습니다.** 임가경제조사·임산물생산비조사 등은 "
                "통계청 마이크로데이터 통합서비스(MDIS, mdis.kostat.go.kr)에서 이용 목적을 "
                "밝히고 받는 자료이고, 이용 약관상 제3자에게 다시 배포할 수 없습니다. "
                "심사위원께서도 같은 경로로 받으실 수 있으며, 받으신 파일을 data/ 아래에 "
                "두고 아래 순서대로 실행하시면 같은 결과가 나옵니다. seed는 42로 "
                "고정했습니다.")

    M.H2(doc, "실행 순서")
    for ln in [
        "pip install -r requirements.txt",
        "",
        "python src/preprocess.py          # 임가경제조사 전처리",
        "python src/preprocess_cost.py     # 임산물생산비조사 전처리 (누수 차단)",
        "python src/train_optuna.py        # Model A 학습 + 3종 벤치마크",
        "python src/train_cost.py          # Model B 학습 + 3종 벤치마크",
        "python src/train_quantile.py      # 예측구간 (P10/P50/P90)",
        "python src/train_panel.py         # 패널 구조 활용 모델",
        "python src/tune_panel.py          # 패널 모델 전용 하이퍼파라미터 탐색",
        "python src/audit_split.py         # 분할 감사",
        "python src/insights.py            # 등급·수령·선도임가",
        "python src/production.py          # 지역 단가·특화도·가공",
        "python src/portfolio.py           # 작목 조합 위험 분산",
        "python src/make_figures.py        # 도표",
        "",
        "uvicorn api.main:app              # 웹 서비스",
    ]:
        code_para(doc, ln, size=8.2, shade="F1F5F2")

    M.H2(doc, "코드 목록")
    rows = [["구분", "파일", "줄 수", "내용"]]
    desc = {
        "src/preprocess.py": "임가경제조사 전처리 · 코드북 파싱 · 누수 변수 차단",
        "src/preprocess_cost.py": "임산물생산비조사 전처리 · 누수 정규식 · 연도별 코드 정합화",
        "src/train_optuna.py": "Model A — Optuna TPE 300회 · 5-fold OOF · 3종 벤치마크",
        "src/train_cost.py": "Model B — Optuna TPE 200회 · 품목별 성능",
        "src/train_quantile.py": "분위수 회귀 · 구간 포함률 검증",
        "src/train_panel.py": "패널 구조 활용 · 임가 단위/연도 단위 분할 비교",
        "src/tune_panel.py": "패널 모델 전용 탐색 (임가 단위 GroupKFold)",
        "src/audit_split.py": "행 단위 분할이 성능을 부풀렸는지 감사",
        "src/explain.py": "TreeSHAP 설명 · 반사실 처방 · 유사 임가 탐색",
        "src/portfolio.py": "작목 조합 위험 분산 · 상관 수축 · PSD 투영",
        "src/insights.py": "등급 전환 · 수령별 수익성 · 선도임가 비교",
        "src/production.py": "지역 단가 · 특화도(LQ) · 가공 손익분기",
        "src/management.py": "임업경영실태조사 — 출하 시기 · 판로",
        "src/subsidy.py": "보조사업 자부담률 기반 실효 ROI",
        "src/region_map.py": "시군구 단위 단가 지도",
        "src/kamis_client.py": "KAMIS 도매가격 API",
        "src/shipping.py": "출하 시기 추천",
        "src/kma_client.py": "기상청 API 허브",
        "src/weather.py": "기상 결합 (시도)",
        "src/weather_sgg.py": "기상 결합 (시군구 230개)",
        "api/main.py": "FastAPI 엔드포인트",
        "api/services.py": "모델 서빙 · 지연 적재",
    }
    total = 0
    for title, files in GROUPS:
        for i, f in enumerate(files):
            path = os.path.join(ROOT, f)
            if not os.path.exists(path):
                continue
            n = len(open(path, encoding="utf-8").read().rstrip("\n").split("\n"))
            total += n
            rows.append([title.split(". ", 1)[1] if i == 0 else "",
                         f, f"{n:,}", desc.get(f, "")])
    rows.append(["**합계**", f"**{len(rows)-1}개 파일**", f"**{total:,}**",
                 "프런트엔드(Vue 3, 4,829줄)는 별도 압축 파일에 있습니다"])
    M.TABLE(doc, rows, widths=[3.4, 4.6, 1.8, 7.2], align_right=[2],
            mark=[len(rows) - 2], size=8.2)

    n_files, n_lines = 0, 0
    for title, files in GROUPS:
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(2)
        h.paragraph_format.keep_with_next = True
        M._font(h.add_run(title), 11, True, M.INK)
        for f in files:
            got = add_code_file(doc, f)
            if got:
                n_files += 1
                n_lines += got
    return n_files, n_lines


def to_pdf(docx_path: str, pdf_path: str) -> bool:
    """만들어진 docx를 HTML로 옮긴 뒤 Chromium으로 인쇄합니다."""
    html_path = os.path.join(ROOT, "docs", "_final_preview.html")
    subprocess.run([sys.executable, "-W", "ignore",
                    os.path.join(ROOT, "src", "preview_docx.py"),
                    docx_path, "-o", html_path], check=True)
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        print("  [건너뜀] playwright가 없어 PDF는 만들지 않았습니다.")
        return False

    with sync_playwright() as p:
        b = p.chromium.launch()
        pg = b.new_page()
        pg.goto("file://" + html_path)
        pg.wait_for_timeout(2500)
        # .page의 그림자와 여백은 화면용입니다. 인쇄에서는 종이가 곧 여백입니다.
        pg.add_style_tag(content="""
            @page { size: A4; margin: 1.7cm 1.9cm 1.5cm; }
            body { background: #fff; }
            .page { width: auto; min-height: 0; margin: 0; padding: 0;
                    box-shadow: none; }
            table, img { break-inside: avoid; }
        """)
        pg.wait_for_timeout(400)
        pg.pdf(path=pdf_path, format="A4", print_background=True,
               margin={"top": "1.7cm", "bottom": "1.5cm",
                       "left": "1.9cm", "right": "1.9cm"})
        b.close()
    os.remove(html_path)
    return True


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--no-pdf", action="store_true")
    a = ap.parse_args()

    if not os.path.exists(BASE):
        sys.exit(f"본문이 없습니다. 먼저 `python src/make_docx.py`를 실행하세요.\n  {BASE}")

    doc = Document(BASE)
    print("[1/2] 본문에 소스코드 별첨을 잇습니다")
    n_files, n_lines = append_appendix(doc)
    doc.save(OUT)
    print(f"  코드 {n_files}개 파일 · {n_lines:,}줄")
    print(f"[saved] {OUT}  ({os.path.getsize(OUT)/1e6:.1f} MB)")

    if not a.no_pdf:
        print("[2/2] PDF")
        if to_pdf(OUT, OUT_PDF):
            print(f"[saved] {OUT_PDF}  ({os.path.getsize(OUT_PDF)/1e6:.1f} MB)")


if __name__ == "__main__":
    main()
