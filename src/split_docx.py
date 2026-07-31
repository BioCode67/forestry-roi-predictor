"""
절별 .docx 나누기 — docs/절별/

한컴독스에 이미 써 넣은 부분이 있으면 문서를 통째로 여는 것이 도움이 되지
않습니다. 웹에서 복사해 붙이는 방법도 표가 자주 어긋납니다. 문서끼리 옮기는
것이 가장 확실합니다. 그래서 절마다 .docx를 따로 만듭니다.

만드는 방법이 조금 특이합니다. 새 문서를 만들어 내용을 옮겨 붙이면 그림이
따라오지 않습니다. 그림은 본문이 아니라 패키지 안의 별도 부품이고, 본문은
관계 아이디(rId)로 그것을 가리킬 뿐이기 때문입니다. 그래서 반대로 합니다.
원본을 통째로 복사한 뒤 필요 없는 문단만 지웁니다. 그림 부품과 관계는 처음부터
그 자리에 있으므로 아이디가 어긋날 일이 없습니다.

쓰지 않는 그림도 파일 안에 남지만 한 파일에 1.5MB 남짓이라 문제가 되지 않습니다.
어긋난 그림을 손으로 다시 넣는 수고보다 낫습니다.

실행: python src/split_docx.py
"""
from __future__ import annotations

import os
import re
import shutil
import zipfile

from docx import Document
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(ROOT, "docs", "제출본_임과함께_데이터분석부문.docx")
OUTDIR = os.path.join(ROOT, "docs", "절별")
ZIP = os.path.join(ROOT, "docs", "제출본_절별.zip")

SECTION_RE = re.compile(r"^([1-5])\)\s*(.+)$")


def section_of(el) -> str | None:
    """이 요소가 절 제목이면 그 번호를 돌려줍니다."""
    if el.tag != qn("w:p"):
        return None
    text = "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()
    m = SECTION_RE.match(text)
    return m.group(1) if m else None


def plan(src: str):
    """원본을 훑어 각 요소가 어느 절에 속하는지 표시합니다."""
    doc = Document(src)
    body = doc.element.body
    marks, cur = [], "0"          # 0 = 표지
    titles = {"0": "표지·요약표"}
    for el in body:
        n = section_of(el)
        if n:
            cur = n
            text = "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()
            titles[n] = text.split(")", 1)[1].strip()
        marks.append(cur)
    return marks, titles


def carve(src: str, dst: str, keep, marks) -> None:
    """원본을 복사한 뒤 남길 절 밖의 요소만 지웁니다.

    keep은 절 번호 하나 또는 여럿입니다. 여럿이면 그 절들이 원래 순서대로
    한 문서에 이어집니다.
    """
    keep = {keep} if isinstance(keep, str) else set(keep)
    shutil.copy2(src, dst)
    doc = Document(dst)
    body = doc.element.body
    # sectPr(용지 설정)은 본문 마지막에 있어야 하므로 건드리지 않습니다
    children = [el for el in body if el.tag != qn("w:sectPr")]
    for el, mark in zip(children, marks):
        if mark not in keep:
            body.remove(el)
    doc.save(dst)


def main() -> None:
    marks, titles = plan(SRC)
    os.makedirs(OUTDIR, exist_ok=True)
    for f in os.listdir(OUTDIR):
        os.remove(os.path.join(OUTDIR, f))

    made = []

    # 한 번에 붙여넣을 묶음 판본. 절을 하나씩 옮기는 것보다 이쪽이 손이 덜 갑니다.
    # 표지는 붙임3 양식에 들어가지 않으므로 뺍니다.
    for name, keep, why in [
        ("붙임3_남은부분_2에서5까지.docx", ["2", "3", "4", "5"],
         "1)절을 이미 넣으셨다면 이 파일 하나면 됩니다"),
        ("붙임3_본문전체_1에서5까지.docx", ["1", "2", "3", "4", "5"],
         "처음부터 다시 넣으실 때"),
    ]:
        dst = os.path.join(OUTDIR, name)
        carve(SRC, dst, keep, marks)
        d = Document(dst)
        made.append((name, len(d.tables), len(d.inline_shapes), os.path.getsize(dst)))
        print(f"  {name:40s} 표 {len(d.tables):2d}개 · 그림 {len(d.inline_shapes):2d}장 · "
              f"{os.path.getsize(dst)/1e6:.1f}MB   ← {why}")
    print()

    for num in ["0", "1", "2", "3", "4", "5"]:
        if num not in titles:
            continue
        label = "0_표지·요약표" if num == "0" else f"{num}_{titles[num]}"
        name = re.sub(r'[\\/:*?"<>|]', "", label).replace(" ", "") + ".docx"
        dst = os.path.join(OUTDIR, name)
        carve(SRC, dst, num, marks)
        d = Document(dst)
        n_img = len(d.inline_shapes)
        made.append((name, len(d.tables), n_img, os.path.getsize(dst)))
        print(f"  {name:40s} 표 {len(d.tables)}개 · 그림 {n_img}장 · "
              f"{os.path.getsize(dst)/1e6:.1f}MB")

    guide = (
        "제출본 파일 — 팀 임과 함께\n\n"
        "★ 가장 빠른 길\n"
        "  1)절까지 넣으셨다면 [붙임3_남은부분_2에서5까지.docx] 하나만 여시면 됩니다.\n"
        "  Ctrl+A → Ctrl+C → 신청서의 2) 자리에 Ctrl+V. 표·그림·글이 한 번에 옵니다.\n"
        "  붙임1 칸은 [붙임1_신청서문구.docx]를 따로 쓰십시오.\n\n"
        "절을 하나씩 옮기고 싶으시면 번호가 붙은 파일들을 쓰시면 됩니다.\n\n"
        "쓰는 법\n"
        "  1. 필요한 절의 .docx를 한컴독스에서 엽니다\n"
        "  2. Ctrl+A 로 전체 선택, Ctrl+C\n"
        "  3. 작업 중인 신청서 문서의 해당 자리에 Ctrl+V\n\n"
        "문서끼리 옮기는 것이라 표·색·그림이 그대로 따라옵니다.\n"
        "웹에서 복사해 붙이는 것보다 확실합니다.\n\n"
        "파일 안에 쓰지 않는 그림이 함께 들어 있습니다. 그림이 어긋나지 않게 하려고\n"
        "원본을 복사한 뒤 필요 없는 부분만 지우는 방식으로 만들었기 때문입니다.\n"
        "화면에는 그 절의 그림만 나옵니다.\n\n"
        "----------------------------------------------------------\n\n"
        + "\n".join(f"{n:44s} 표 {t}개 · 그림 {i}장" for n, t, i, _ in made) + "\n")
    with open(os.path.join(OUTDIR, "00_읽어보기.txt"), "w", encoding="utf-8") as f:
        f.write(guide)

    with zipfile.ZipFile(ZIP, "w", zipfile.ZIP_DEFLATED) as z:
        for f in sorted(os.listdir(OUTDIR)):
            z.write(os.path.join(OUTDIR, f), f)
    print(f"\n[saved] {ZIP}  ({os.path.getsize(ZIP)/1e6:.1f} MB)")


if __name__ == "__main__":
    main()
