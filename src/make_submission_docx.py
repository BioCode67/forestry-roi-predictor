"""
요강 양식에 맞춘 제출 서류 일체 — docs/제출서류_일체_임과함께.docx (+ .pdf)

forestry_download.pdf의 붙임 양식을 그대로 옮겨 채웁니다. 순서도 요강과 같게
붙임1 → 붙임3 → 붙임4 → 붙임5 → 붙임6 → 별첨으로 둡니다.

붙임3 본문은 make_docx.py가 만든 문서에서 그대로 가져옵니다. 여기서 다시 짜면
두 벌이 생기고, 한쪽만 고치는 일이 반드시 따라옵니다.

서명란은 비워 둡니다. 이름은 채우되 (인)·(서명) 자리는 손으로 하실 몫입니다.
연락처와 전자우편도 확인이 필요해 표시만 해 둡니다.

실행: python src/make_submission_docx.py [--no-pdf]
"""
from __future__ import annotations

import argparse
import os
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import make_docx as M
from make_final_docx import append_appendix, to_pdf

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
BODY = os.path.join(ROOT, "docs", "제출본_임과함께_데이터분석부문.docx")
OUT = os.path.join(ROOT, "docs", "제출서류_일체_임과함께.docx")
OUT_PDF = os.path.join(ROOT, "docs", "제출서류_일체_임과함께.pdf")

# 신청인이 확인해 채워야 하는 자리. 눈에 띄게 표시해 둡니다.
BLANK = "＿＿＿＿＿＿＿＿"

APPLICANT = {"성명": "김주형", "소속": "영남대학교"}
TEAM = "임과 함께"
TITLE = "임업통계 마이크로데이터 기반 임가 맞춤형 수익성 예측 및 경영 개선 시스템"
TODAY = ("2026", "7", "31")

SH_LABEL = "EDF1F7"   # 양식 왼쪽 항목 칸 — 원본의 옅은 보라 계열
SH_HEAD = "F3F1F8"


def form_cell(cell, text, bold=False, size=9.6, center=False, fill=None,
              align_left_pad=True):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    M._rich(p, text, size, M.INK)
    if bold:
        for r in p.runs:
            r.font.bold = True
    if fill:
        M._shade(cell, fill)
    return p


def form_table(doc, rows, widths, height=None):
    t = doc.add_table(rows=rows, cols=len(widths))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, w in enumerate(widths):
        for row in t.rows:
            row.cells[ci].width = Cm(w)
    if height:
        for row in t.rows:
            row.height = Cm(height)
    return t


def sheet_title(doc, tag, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    M._font(p.add_run(tag), 10, True, RGBColor(0x2F, 0x4F, 0x8F))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    M._font(p.add_run(title), 17, True, M.INK)
    M._border(p, "bottom", "9C8FC4", 12)


def sign_block(doc, lines, name_label="신청인(대표자)", mark="(인)"):
    for ln in lines:
        M.P(doc, ln, 10.3, after=3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(10)
    M._font(p.add_run(f"{TODAY[0]}년   {TODAY[1]}월   {TODAY[2]}일"), 11, False, M.INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(4)
    M._font(p.add_run(f"{name_label}   {APPLICANT['성명']}        {mark}"), 11, False, M.INK)


# ── 붙임1 참가 신청서 ─────────────────────────────────────────────────────
def sheet_form1(doc, overview: str, summary: str, data_rows):
    sheet_title(doc, "[붙임 1] 공통 양식(신청서)",
                "「2026년 임업통계 활용 경진대회」참가 신청서")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(3)
    M._font(p.add_run("* 해당란에 ☑ 표시"), 8.5, False, M.MUTED)

    # 위쪽 — 분야·구분·기획명
    t = form_table(doc, 4, [3.2, 6.4, 6.8])
    for ri, (label, a, b) in enumerate([
        ("공모 분야", "□ 아이디어", "☑ 데이터 분석"),
        ("신청자 정보", "☑ 학생", "□ 일반"),
        ("참가 구분", "☑ 개인", "□ 팀"),
    ]):
        form_cell(t.cell(ri, 0), label, bold=True, center=True, fill=SH_LABEL)
        form_cell(t.cell(ri, 1), a)
        form_cell(t.cell(ri, 2), b)
    form_cell(t.cell(3, 0), "팀 명 /\n기획명(주제)", bold=True, center=True, fill=SH_LABEL)
    form_cell(t.cell(3, 1), TEAM, center=True)
    form_cell(t.cell(3, 2), TITLE)

    # 기획 개요 · 요약
    t = form_table(doc, 2, [3.2, 13.2])
    form_cell(t.cell(0, 0), "기획 개요\n(분석·활용 개요)", bold=True, center=True, fill=SH_LABEL)
    form_cell(t.cell(0, 1), overview, size=9.0)
    form_cell(t.cell(1, 0), "기획서 요약", bold=True, center=True, fill=SH_LABEL)
    form_cell(t.cell(1, 1), summary, size=9.0)

    # 활용데이터 정보
    t = form_table(doc, 1 + len(data_rows), [3.2, 3.4, 3.6, 6.2])
    form_cell(t.cell(0, 0), "활용데이터 정보\n※복수기재가능,\n임업통계 정보 필수",
              bold=True, center=True, size=8.6, fill=SH_LABEL)
    for ci, h in enumerate(["출처", "제공기관명", "데이터명"], start=1):
        form_cell(t.cell(0, ci), h, bold=True, center=True, fill=SH_HEAD)
    for ri, (src, org, name) in enumerate(data_rows, start=1):
        form_cell(t.cell(ri, 0), "", center=True)
        form_cell(t.cell(ri, 1), f"{ri}. {src}", size=9.0)
        form_cell(t.cell(ri, 2), org, size=9.0)
        form_cell(t.cell(ri, 3), name, size=9.0)
    # 왼쪽 항목 칸을 세로로 합칩니다
    t.cell(0, 0).merge(t.cell(len(data_rows), 0))

    # 참가자 정보
    t = form_table(doc, 3, [3.2, 3.0, 3.0, 3.4, 3.8])
    form_cell(t.cell(0, 0), "참가자 정보", bold=True, center=True, fill=SH_LABEL)
    for ci, h in enumerate(["성 명", "소 속", "연락처(휴대전화)", "전자우편"], start=1):
        form_cell(t.cell(0, ci), h, bold=True, center=True, fill=SH_HEAD)
    form_cell(t.cell(1, 1), APPLICANT["성명"], center=True)
    form_cell(t.cell(1, 2), APPLICANT["소속"], center=True)
    form_cell(t.cell(1, 3), BLANK, center=True, size=9.0)
    form_cell(t.cell(1, 4), BLANK, center=True, size=9.0)
    for ci in range(1, 5):
        form_cell(t.cell(2, ci), "", center=True)
    t.cell(0, 0).merge(t.cell(2, 0))

    # 이전 수혜 이력
    t = form_table(doc, 2, [3.2, 2.6, 10.6])
    form_cell(t.cell(0, 0), "이전 수혜 이력\n및 입상 실적", bold=True, center=True, fill=SH_LABEL)
    form_cell(t.cell(0, 1), "년도", bold=True, center=True, fill=SH_HEAD)
    form_cell(t.cell(0, 2), "내용", bold=True, center=True, fill=SH_HEAD)
    form_cell(t.cell(1, 1), "—", center=True)
    form_cell(t.cell(1, 2), "해당 없음", center=True)
    t.cell(0, 0).merge(t.cell(1, 0))

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    sign_block(doc, [
        "본인(팀)은 ‘2026년 임업통계 활용 경진대회’ 참가와 관련하여 제출한 사항에 허위가 "
        "없으며, 유의사항을 숙지하고 진행에 필요한 사항에 성실히 응할 것을 동의합니다.",
    ])


# ── 붙임4 참가 서약서 ────────────────────────────────────────────────────
def sheet_form4(doc):
    doc.add_page_break()
    sheet_title(doc, "[붙임 4] 공통 양식(참가 서약서)", "참 가 서 약 서")

    M.P(doc, "본인(팀)은 「2026년 임업통계 활용 경진대회」에 출품하며 아래 사항을 숙지하고, "
             "허위사실 기재 및 타인의 권리를 침해하는 등의 행위로 인하여 손해를 발생시키는 "
             "경우, 본인의 귀책으로 인하여 발생되는 손해에 관한 손해배상책임이 본인에게 "
             "있음을 확인합니다.", 10.3, after=10)
    for n, txt in enumerate([
        "이미 채택된 제안과 동일한 것, 표절 및 복제 등의 지식재산권 침해 작품, 타 경진대회 "
        "입상작품 등은 심사에서 제외되며, 이에 따른 모든 법적 책임은 참가자에게 있음",
        "제출한 작품이 제3자의 권리(소유권, 저작권, 이용권)를 침해하였거나 이와 관련한 "
        "분쟁이 발생한 사실이 없으며, 이로 인하여 발생하는 법적인 책임은 출품자에게 있음",
        "수상 이후 위반 사실이 밝혀질 경우 수상 취소 및 상금 환수(자진반납)에 이의를 "
        "제기하지 않음",
    ], start=1):
        p = M.P(doc, f"{n}. {txt}", 10.3, after=6)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
    M.P(doc, "본인은 유의사항을 충분히 숙지하였으며 대회진행에 필요한 주관기관의 요구사항에 "
             "성실히 응할 것에 동의합니다.", 10.3, after=4)
    sign_block(doc, [], name_label="서약자 : 성명", mark="(서명 또는 인)")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    M._font(p.add_run("한국임업진흥원장, 한국산림복지진흥원장  귀중"), 11, True, M.INK)


# ── 붙임5 개인정보 동의서 ────────────────────────────────────────────────
def sheet_form5(doc):
    doc.add_page_break()
    sheet_title(doc, "[붙임 5] 공통 양식(개인정보 수집·이용·활용 동의서)",
                "개인정보 수집·이용 동의서")

    M.P(doc, "■ 개인정보 수집·이용 동의", 10.5, M.FOREST, after=4)
    M.P(doc, "한국임업진흥원과 한국산림복지진흥원은 2026년 임업통계 활용 경진대회(이하 "
             "경진대회)를 위하여 아래와 같이 개인정보를 수집·이용합니다.", 9.8, M.MUTED, after=6)
    M.TABLE(doc, [
        ["수집목적", "수집항목", "보유기간"],
        ["· 경진대회 심사의 진행\n· 진행단계별 결과 등 관련 내용 공지\n"
         "· 후속지원을 위한 수요조사\n· 신규 개방 공공데이터 수요조사\n· 포상금 지급",
         "(필수) 팀명(성명), 생년월일, 주소, 전화번호, 휴대폰번호, E-Mail, 소속, "
         "포상금 수령 계좌정보\n(선택) 팀장성명, 팩스번호, 팀원성명, 팀원소속, "
         "팀원 연락처, 포상금 수령 계좌정보",
         "경진대회 및\n후속지원 종료 후\n3개월"],
    ], widths=[5.0, 8.4, 3.0], size=8.6)
    M.P(doc, "■ 동의를 거부할 권리 및 동의 거부에 따른 불이익", 10.5, M.FOREST, after=4)
    M.P(doc, "귀하는 개인정보 수집·이용에 동의하지 않을 권리가 있습니다. 다만 개인정보를 "
             "제공받지 못할 경우 경진대회 심사진행이 어려울 수 있으며 따라서 개인정보 제공에 "
             "동의하지 않은 경우 경진대회 참가가 제한될 수 있습니다.", 9.6, M.MUTED, after=6)
    M.P(doc, "☞ 위 개인정보 수집·이용에 동의하십니까?      **☑ 동의**      □ 미동의",
        10.3, after=12)

    M.P(doc, "■ 개인정보 제3자 제공에 대한 별도 동의", 10.5, M.FOREST, after=4)
    M.TABLE(doc, [
        ["개인정보를 제공받는 자", "이용 목적", "제공하는 항목", "보유기간"],
        ["· 프로퍼커뮤니케이션",
         "· 경진대회 접수·심사 등 진행\n· 진행단계별 결과 등 관련 내용 공지\n"
         "· 후속지원을 위한 수요조사 및 후속지원",
         "팀명(성명), 생년월일, 주소, 전화번호, 휴대폰번호, E-Mail, 소속, 팀장성명, "
         "팩스번호, 팀원성명, 팀원소속, 팀원 연락처, 포상금 수령 계좌정보",
         "경진대회\n종료 후 3개월"],
    ], widths=[3.4, 5.0, 5.4, 2.6], size=8.6)
    M.P(doc, "귀하는 개인정보 제3자 제공에 동의하지 않을 권리가 있습니다. 제3자에게 제공하는 "
             "정보는 경진대회 평가에 필수 항목으로 해당정보를 제공하지 못할 경우 심사진행이 "
             "어려울 수 있습니다.", 9.6, M.MUTED, after=6)
    M.P(doc, "☞ 위와 같이 개인정보를 제3자에 제공하는데 동의하십니까?      **☑ 동의**      "
             "□ 미동의", 10.3, after=6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    M._font(p.add_run(f"{TODAY[0]}년   {TODAY[1]}월   {TODAY[2]}일"), 11, False, M.INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    M._font(p.add_run(f"신청인   소속 {APPLICANT['소속']}   성명 {APPLICANT['성명']}"
                      "        (서명)"), 11, False, M.INK)


# ── 붙임6 출품작 공개·공유 동의서 ────────────────────────────────────────
def sheet_form6(doc):
    doc.add_page_break()
    sheet_title(doc, "[붙임 6] 공통양식(동의서)", "출품작 제3자 공개·공유 동의서")

    M.P(doc, "한국임업진흥원은 「2026년 임업통계 활용 경진대회」 진행에 따른 출품작에 대하여 "
             "다음과 같이 귀하의 동의를 얻고자 합니다.", 10.3, after=10)
    for n, head, body in [
        (1, "출품작 제3자 공개·공유 목적",
         "응모된 출품작에 대한 평가와 경진대회 관리 및 운영에 관련한 업무수행을 위함"),
        (2, "출품작 공개·공유 항목",
         "응모된 출품작의 목적, 사용하는 데이터 종류, 효과성 등"),
        (3, "출품작 보유·이용기간",
         "수집된 개인정보는 경진대회 결과 최종발표일로부터 2년 이내에 폐기하며, 수상자의 "
         "사후관리나 중앙의 경진대회에 추천할 경우 본인의 동의를 얻어 출품작을 관련 기관에 "
         "제공할 수 있음"),
    ]:
        M.P(doc, f"{n}. **{head}**", 10.3, after=2)
        p = M.P(doc, f"⦁ {body}", 10.0, M.MUTED, after=8)
        p.paragraph_format.left_indent = Cm(0.6)

    M.P(doc, "출품작에 대한 동의 여부      **☑ 동의함**      □ 동의하지 않음",
        10.5, after=10)
    M.P(doc, "※ 동의를 거부할 권리와 거부에 따른 불이익 — 지원자는 제출한 출품작의 "
             "공개·공유를 거부할 권리가 있습니다. 다만, 지원자가 동의를 거부하는 경우 "
             "심사대상에서 제외될 수 있음을 알려드립니다.", 9.6, M.MUTED, after=4)
    M.P(doc, "※ 기타 자세한 사항은 ‘2026년 임업통계 활용 경진대회 운영 사무국’으로 "
             "문의바랍니다.", 9.6, M.MUTED)


# ── 본문 다루기 ──────────────────────────────────────────────────────────
def open_body_copy() -> Document:
    """붙임3 본문 문서를 복사해 엽니다.

    새 문서에 요소만 옮기면 그림이 따라오지 않습니다. 그림은 본문이 아니라
    패키지 안의 별도 부품이고 본문은 관계 아이디로 그것을 가리킬 뿐이기
    때문입니다. 복사본에서 시작하면 부품과 관계가 처음부터 제자리에 있습니다.
    """
    import re
    import shutil

    shutil.copy2(BODY, OUT)
    doc = Document(OUT)
    body = doc.element.body

    # 표지는 양식에 들어가지 않습니다. 첫 절 제목 앞을 걷어냅니다.
    dropped = 0
    for el in list(body):
        if el.tag == qn("w:sectPr"):
            continue
        if el.tag == qn("w:p"):
            text = "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()
            if re.match(r"^1\)\s", text):
                break
        body.remove(el)
        dropped += 1
    return doc, dropped


def move_to_front(doc, n_new: int) -> None:
    """마지막에 만든 요소 n_new개를 문서 맨 앞으로 옮깁니다.

    python-docx는 뒤에만 붙일 수 있습니다. 붙임1을 앞에 두려면 만든 뒤에
    옮기는 수밖에 없습니다.
    """
    body = doc.element.body
    kids = [el for el in body if el.tag != qn("w:sectPr")]
    for i, el in enumerate(kids[-n_new:]):
        body.remove(el)
        body.insert(i, el)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--no-pdf", action="store_true")
    a = ap.parse_args()

    if not os.path.exists(BODY):
        sys.exit(f"본문이 없습니다. 먼저 `python src/make_docx.py`를 실행하세요.\n  {BODY}")

    from make_paste_page import FORM_HTML
    import re as _re
    blocks = [_re.sub(r"<[^>]+>", "", m.group(2)).strip()
              for m in _re.finditer(r"<p([^>]*)>(.*?)</p>", FORM_HTML, _re.S)]
    heads = [i for i, m in enumerate(_re.finditer(r"<p([^>]*)>", FORM_HTML))
             if "11pt" in m.group(1)]
    overview = "\n".join(b for b in blocks[heads[0] + 1:heads[1]] if b)
    summary = "\n".join(b for b in blocks[heads[1] + 1:] if b)

    data_rows = [
        ("통계청 MDIS", "산림청 · 한국임업진흥원", "임가경제조사 마이크로데이터 (2019~2023)"),
        ("통계청 MDIS", "산림청 · 한국임업진흥원", "임산물생산비조사 마이크로데이터 (2018~2024)"),
        ("통계청 MDIS", "산림청 · 한국임업진흥원", "임산물생산조사 마이크로데이터 (2022~2024)"),
        ("통계청 MDIS", "산림청 · 한국임업진흥원", "임업경영실태조사 마이크로데이터 (2018·2020)"),
        ("KAMIS", "한국농수산식품유통공사(aT)", "농수산물 도매시장 가격 정보"),
        ("기상청 API 허브", "기상청", "종관기상관측(ASOS) 일자료"),
        ("산림청 누리집", "산림청", "산림사업 보조금 지원 사업 목록"),
    ]

    print("[1/4] 붙임3 본문 복사 (그림·서식 그대로)")
    doc, dropped = open_body_copy()
    print(f"  표지 {dropped}개 요소 제거")

    print("[2/4] 붙임1 참가 신청서")
    before = len([el for el in doc.element.body if el.tag != qn("w:sectPr")])
    sheet_form1(doc, overview, summary, data_rows)

    # 붙임3 표제도 붙임1 뒤, 본문 앞에 들어가야 합니다
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    M._font(p.add_run("[붙임 3] 데이터 분석 부문 양식"), 10, True,
            RGBColor(0x2F, 0x4F, 0x8F))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    M._font(p.add_run("「2026년 임업통계 활용 경진대회」양식(데이터 분석)"), 15, True, M.INK)
    M._border(p, "bottom", "9C8FC4", 12)

    after = len([el for el in doc.element.body if el.tag != qn("w:sectPr")])
    move_to_front(doc, after - before)
    print(f"  요소 {after - before}개를 맨 앞으로")

    print("[3/4] 붙임4·5·6")
    sheet_form4(doc)
    sheet_form5(doc)
    sheet_form6(doc)

    print("[4/4] 별첨 — 분석 코드")
    n_files, n_lines = append_appendix(doc)
    print(f"  코드 {n_files}개 파일 · {n_lines:,}줄")

    doc.save(OUT)
    print(f"[saved] {OUT}  ({os.path.getsize(OUT)/1e6:.1f} MB)")

    if not a.no_pdf:
        if to_pdf(OUT, OUT_PDF):
            print(f"[saved] {OUT_PDF}  ({os.path.getsize(OUT_PDF)/1e6:.1f} MB)")


if __name__ == "__main__":
    main()
