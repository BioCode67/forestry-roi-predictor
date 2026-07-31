"""
붙임1(신청서) 문구 .docx — docs/절별/붙임1_신청서문구.docx

붙임3 본문이 아니라 양식 칸을 채우는 글이라 제출본 문서에서 뽑아 올 수 없습니다.
문구는 make_paste_page.py의 FORM_HTML 한 곳에만 두고, 여기서는 그것을 읽어
문서로 옮깁니다. 두 곳에 따로 적으면 한쪽만 고치고 넘어가게 됩니다.

실행: python src/make_form_docx.py
"""
from __future__ import annotations

import os
import re

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from make_docx import FONT, FOREST, H2, P, _font
from make_paste_page import FORM_HTML

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "docs", "절별", "붙임1_신청서문구.docx")

TAG = re.compile(r"<[^>]+>")


def blocks():
    """FORM_HTML을 (제목여부, 글) 목록으로 풉니다."""
    for m in re.finditer(r"<p([^>]*)>(.*?)</p>", FORM_HTML, re.S):
        attrs, inner = m.group(1), m.group(2)
        # <b>는 **로 바꿔 P()가 굵게 처리하게 둡니다
        text = re.sub(r"</?b>", "**", inner)
        text = TAG.sub("", text)
        text = re.sub(r"\s+", " ", text).strip()
        if text:
            yield ("11pt" in attrs, text)


def main() -> None:
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.0)
    sec.top_margin = sec.bottom_margin = Cm(1.8)
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    _font(p.add_run("[붙임1] 참가 신청서 — 붙여넣을 문구"), 13, True, FOREST)

    for is_head, text in blocks():
        if is_head:
            H2(doc, text)
        else:
            P(doc, text)

    doc.save(OUT)
    print(f"[saved] {OUT}")
    print(f"  문단 {len(doc.paragraphs)}개")


if __name__ == "__main__":
    main()
