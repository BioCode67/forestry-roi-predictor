"""
제출본 .docx 생성 — docs/제출본_임과함께_데이터분석부문.docx

한컴독스(HWP)로 옮겨 편집할 것을 전제로 만듭니다. 마크다운은 표와 그림이
붙여넣기에서 깨지므로, 표는 실제 Word 표로, 그림은 문서에 삽입된 이미지로 넣습니다.

글꼴은 맑은 고딕으로 두되 동아시아 글꼴 속성을 따로 지정합니다. 이 속성을 빼면
한컴독스에서 한글만 다른 글꼴로 튑니다.

실행: python src/make_docx.py
"""
from __future__ import annotations

import json
import os
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG = os.path.join(ROOT, "docs", "figures")
OUT = os.path.join(ROOT, "docs", "제출본_임과함께_데이터분석부문.docx")

FONT = "맑은 고딕"
INK = RGBColor(0x1A, 0x1E, 0x21)
MUTED = RGBColor(0x55, 0x5F, 0x66)
FOREST = RGBColor(0x1F, 0x5A, 0x38)
AMBER = RGBColor(0x8A, 0x53, 0x00)

SH_HEAD = "E8EFEA"     # 표 머리
SH_ALT = "F7F9F8"      # 표 짝수 줄
SH_NOTE = "F2F7F4"     # 강조 상자
SH_WARN = "FDF6EA"


# ── 저수준 헬퍼 ────────────────────────────────────────────────────────────
def _font(run, size=10.5, bold=False, color=INK, name=FONT, ea=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    # 동아시아 글꼴을 따로 지정하지 않으면 한글만 다른 글꼴로 떨어진다
    run._element.rPr.rFonts.set(qn("w:eastAsia"), ea or name)


def _shade(cell_or_par, hex_color):
    el = cell_or_par._tc if hasattr(cell_or_par, "_tc") else cell_or_par._p.get_or_add_pPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_color)
    el.append(sh)


def _border(par, side="left", color="2E7D4F", sz=18):
    pPr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    b = OxmlElement(f"w:{side}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(sz))
    b.set(qn("w:space"), "8")
    b.set(qn("w:color"), color)
    pbdr.append(b)
    pPr.append(pbdr)


BOLD = re.compile(r"\*\*(.+?)\*\*")


def _rich(par, text, size=10.5, color=INK):
    """**굵게** 표기만 해석합니다. 그 이상은 문서를 어지럽힙니다."""
    pos = 0
    for m in BOLD.finditer(text):
        if m.start() > pos:
            _font(par.add_run(text[pos:m.start()]), size, False, color)
        _font(par.add_run(m.group(1)), size, True, color)
        pos = m.end()
    if pos < len(text):
        _font(par.add_run(text[pos:]), size, False, color)


# ── 문단 API ──────────────────────────────────────────────────────────────
def H1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True
    _font(p.add_run(text), 15, True, FOREST)
    _border(p, "bottom", "C8D8CE", 8)


def H2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    _font(p.add_run(text), 11.5, True, INK)


def P(doc, text, size=10.5, color=INK, indent=0, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.32
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    _rich(p, text, size, color)
    return p


def BUL(doc, text, size=10.5, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.32
    p.paragraph_format.left_indent = Cm(0.7 + 0.6 * level)
    _rich(p, text, size)
    return p


def NOTE(doc, text, kind="good"):
    """본문에서 눈에 걸리게 할 대목. 배경색과 왼쪽 선을 함께 줍니다."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.right_indent = Cm(0.2)
    p.paragraph_format.line_spacing = 1.4
    _rich(p, text, 10)
    _border(p, "left", "2E7D4F" if kind == "good" else "D97706", 20)  # pBdr가 shd보다 앞
    _shade(p, SH_NOTE if kind == "good" else SH_WARN)
    return p


def CODE(doc, text):
    for line in text.strip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing = 1.12
        _font(p.add_run(line or " "), 8.6, False, RGBColor(0x2B, 0x35, 0x3A), "Consolas", FONT)
        _shade(p, "F4F6F7")


def TABLE(doc, rows, widths=None, align_right=None, size=9.1):
    """첫 줄을 머리로 보고 음영을 넣습니다. align_right는 우측 정렬할 열 번호."""
    align_right = align_right or []
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0.8)
            p.paragraph_format.space_before = Pt(0.8)
            p.paragraph_format.line_spacing = 1.12
            if ci in align_right and ri > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif ri == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _rich(p, str(val), size, INK)
            if ri == 0:
                for r in p.runs:
                    r.font.bold = True
                _shade(cell, SH_HEAD)
            elif ri % 2 == 0:
                _shade(cell, SH_ALT)
    if widths:
        for ci, w in enumerate(widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def FIGURE(doc, name, caption, width=16.4):
    path = os.path.join(FIG, name)
    if not os.path.exists(path):
        print("  [warn] 그림 없음:", name)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Cm(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    _font(c.add_run(caption), 9, False, MUTED)


# ── 문서 ──────────────────────────────────────────────────────────────────
def build():
    m = {f: json.load(open(os.path.join(ROOT, "models", f), encoding="utf-8"))
         for f in ("metrics_summary.json", "metrics_cost.json",
                   "metrics_quantile.json", "portfolio.json", "insights.json")}
    A, B = m["metrics_summary.json"], m["metrics_cost.json"]
    Q, PF, IN = m["metrics_quantile.json"], m["portfolio.json"], m["insights.json"]

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(1.9)
    sec.top_margin = Cm(1.7)
    sec.bottom_margin = Cm(1.5)

    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    # ── 표지 머리 ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    _font(p.add_run("2026년 임업통계 활용 경진대회  |  데이터 분석 부문"), 10.5, False, MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    _font(p.add_run("임가별 수익률 예측과 처방을 제공하는\n임업 의사결정 지원 시스템"),
          19, True, FOREST)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    _font(p.add_run("— 임업통계 마이크로데이터 4종 + 공공데이터 3종 융복합 —"), 10.5, False, MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    _font(p.add_run("팀명 : 임과 함께"), 12.5, True, INK)

    TABLE(doc, [
        ["구분", "내용"],
        ["공모 분야", "데이터 분석"],
        ["활용 통계 (필수)", "임가경제조사 · 임산물생산비조사 · 임산물생산조사 · 임업경영실태조사"],
        ["융복합 공공데이터", "KAMIS 도매가격(aT) · 기상청 ASOS · 산림사업 보조금 목록 · 행정구역 경계"],
        ["분석 표본", f"Model A {A['dataset']['rows']:,}건(설명변수 {A['dataset']['n_features']}개) · "
                    f"Model B {B['dataset']['rows']:,}건(설명변수 {B['dataset']['n_features']}개)"],
        ["핵심 성과", "집단 평균 대비 설명력 4.05배 / 5.95배 (패널 구조 활용 시 10.2배) · "
                  "예측구간 포함률 78.2% / 76.5% 검증 · 데이터 누수 자체 적발 및 교정"],
        ["산출물", "분석 코드 전체(별첨) · FastAPI + Vue 3 웹 애플리케이션"],
    ], widths=[3.6, 12.8])

    # ══════════════════════════════════════════════════════════ 1)
    H1(doc, "1) 추진배경 및 필요성")

    H2(doc, "가. 현황 — 통계는 쌓이는데, 임가는 자기 숫자를 모른다")
    P(doc, "산림청은 임가경제조사, 임산물생산비조사, 임산물생산조사, 임업경영실태조사 등 "
           "국가승인통계를 매년 생산하고 통계청 MDIS를 통해 마이크로데이터까지 공개하고 있습니다. "
           "그러나 이 통계가 임가에게 되돌아오는 형태는 대부분 **집단 평균**입니다. "
           "“밤 재배 임가의 평균 임업소득은 얼마”와 같은 문장이 그것입니다.")
    P(doc, "문제는 임업 수익의 흩어짐이 지나치게 크다는 데 있습니다. 본 분석에서 임가경제조사 "
           f"{A['dataset']['rows']:,}개 임가-연도 표본의 ROI(임업소득÷임업경영비)는 "
           f"**평균 {A['dataset']['roi_train_mean']:.1f}%, 표준편차 {A['dataset']['roi_train_std']:.1f}%p**"
           "였습니다. 표준편차가 평균보다 큽니다. 이런 분포에서 평균은 대표값 구실을 하지 못합니다.")
    NOTE(doc, "실제로 **지역별×업종별 집단 평균을 그대로 예측값으로 쓰면 Test R²는 "
              f"{A['forest_service_baseline']['test']['R2']:.4f}**에 그칩니다. 현장에서 널리 쓰이는 "
              "“같은 지역, 같은 품목의 평균을 참고하라”는 방식이 개별 임가의 수익 변동을 "
              "사실상 4%밖에 설명하지 못한다는 뜻입니다.", "warn")

    H2(doc, "나. 문제점 — 세 가지 공백")
    TABLE(doc, [
        ["공백", "내용", "임가가 겪는 일"],
        ["① 개인화 부재", "집단 평균만 제공", "“우리 산은 저 평균과 다른데”"],
        ["② 불확실성 은폐", "단일 숫자만 제시", "평균을 확정값으로 오해 → 과잉 투자"],
        ["③ 처방 부재", "현황 기술에 그침", "“그래서 무엇을 바꾸라는 것인가”"],
    ], widths=[3.4, 5.4, 7.6])
    P(doc, "특히 ②는 위험합니다. ROI 216%p의 산포를 감춘 채 “평균 155%”만 전달하면, "
           "임가는 그 수치를 기대값이 아니라 보장값으로 읽습니다.")

    H2(doc, "다. 필요성")
    P(doc, "따라서 필요한 것은 새로운 통계가 아니라 **이미 있는 마이크로데이터를 임가 단위로 "
           "되돌려주는 분석 계층**입니다. 구체적으로 세 가지를 갖춰야 합니다.")
    BUL(doc, "**개인화** — 임가의 조건(지역·업종·규모·경영비·연령 등)을 입력하면 그 임가의 예측치를 산출")
    BUL(doc, "**불확실성 제시** — 점추정이 아니라 구간(P10~P90)으로 제시하고, 그 구간이 실제로 맞는지 검증")
    BUL(doc, "**처방** — “무엇을 얼마나 바꾸면 수익이 어떻게 달라지는가”를 반사실(counterfactual)로 탐색")
    P(doc, "본 과제는 이 세 가지를 실제로 구현하고, **전국 어느 임가든 웹에서 바로 쓸 수 있는 "
           "서비스**까지 완성했습니다.")

    FIGURE(doc, "fig09_distribution.png",
           "[그림 1] 임가 ROI 분포 — 평균은 대표값이 아니다", 12.1)
    P(doc, "분포가 오른쪽으로 길게 늘어져 있어 평균(155.5%)이 중앙값(약 87%)보다 크게 위쪽에 "
           "놓입니다. 임가의 80%는 −38%에서 429% 사이에 흩어져 있습니다. 이 폭을 감춘 채 하나의 "
           "평균만 전달하는 것이 현행 통계 환류의 가장 큰 문제입니다.", 9.8, MUTED)
    FIGURE(doc, "fig08_data.png", "[그림 2] 활용 데이터 구성 — 임업통계 4종(필수) + 공공데이터 3종", 11.1)

    # ══════════════════════════════════════════════════════════ 2)
    H1(doc, "2) 아이디어 기획 세부 내용")

    H2(doc, "가. 개요")
    P(doc, "**임업통계 마이크로데이터로 학습한 두 개의 예측 모델과, 그 예측을 임가가 쓸 수 있는 "
           "행동으로 바꾸는 분석 계층, 그리고 이를 담은 웹 애플리케이션**입니다.")
    TABLE(doc, [
        ["단계", "무엇을 하는가", "쓰인 기법"],
        ["① 예측", "얼마나 남을지 구간과 함께 제시", "XGBoost + 분위수 회귀"],
        ["② 설명", "왜 그 숫자가 나왔는지 분해", "TreeSHAP"],
        ["③ 처방", "무엇을 바꾸면 되는지 탐색", "반사실(counterfactual) 격자 탐색"],
        ["④ 비교", "비슷한 임가는 어떤지 제시", "최근접 이웃 (로그 척도)"],
        ["⑤ 전략", "언제 팔지 · 무엇을 섞을지", "KAMIS 시세 · 평균–분산 최적화"],
    ], widths=[2.4, 7.4, 6.6])

    H2(doc, "나. 주요 기능")
    P(doc, "**기능 1 — 임가별 수익률 예측 (Model A).** 임가경제조사 2019~2023년 마이크로데이터 "
           f"{A['dataset']['rows']:,}개 표본, 설명변수 {A['dataset']['n_features']}개. "
           "지역·업종·임지규모·전겸업·연령·가구원수 등 임가 제원과 임업경영비·자본을 입력받아 "
           f"ROI를 예측합니다. Test R² {A['optuna_xgboost']['test']['R2']:.4f} — "
           "산림청 방식(집단 평균) 대비 **4.05배**.")
    P(doc, "**기능 2 — 품목별 정밀 진단 (Model B).** 임산물생산비조사 2018~2024년 "
           f"{B['dataset']['rows']:,}개 관측, 설명변수 {B['dataset']['n_features']}개, 5개 품목"
           "(밤·대추·떫은감·표고 노지·표고 톱밥). 비목별 투입 구조를 입력받아 단위면적당 ROI를 "
           f"예측합니다. Test R² {B['optuna_xgboost']['test']['R2']:.4f} — 베이스라인 대비 **5.95배**.")
    P(doc, "**기능 3 — 예측구간 제시 (분위수 회귀).** reg:quantileerror 목적함수로 P10/P50/P90을 "
           "별도 학습했습니다. 실측값이 P10~P90 안에 들어간 비율은 "
           f"Model A {Q['roi']['coverage_80pct']*100:.1f}%, Model B {Q['cost']['coverage_80pct']*100:.1f}%로 "
           "**명목 80%에 근접**합니다. 즉 구간 자체가 검증된 값입니다.")
    P(doc, "**기능 4 — 설명가능 AI (TreeSHAP).** pred_contribs로 예측을 변수별 기여로 가법 분해합니다. "
           "파생변수(임업경영비 / log_임업경영비 / ha당_경영비)는 사람이 읽을 수 있도록 하나로 묶어 "
           "“예측 ROI가 평균보다 40%p 낮은 것은 경영비 규모가 −25%p, 지역이 −10%p…” 형태로 제시합니다.")
    P(doc, "**기능 5 — 반사실 처방 (Counterfactual).** 바꿀 수 있는 변수만 골라 격자 탐색합니다.")
    NOTE(doc, "실행 가능성을 기준으로 변수를 걸렀습니다. **임지 규모는 제외했습니다.** "
              "“산을 1ha 미만으로 줄이면 수익률이 오른다”는 결과는 통계적으로 참일 수 있으나 "
              "땅을 팔라는 말이라 조언이 될 수 없습니다. 남긴 변수는 임업경영비(올해 실행 가능), "
              "전/겸업 전환(중기), 작목 전환(장기)입니다. 목표치도 처음에는 동일 조건 상위 25%로 "
              "두었으나 322%처럼 도달 불가능한 값이 나와 **동일 조건 임가의 중앙값**으로 바꾸었습니다.")
    P(doc, "**기능 6 — 유사 임가 탐색.** 로그 척도 거리로 조건이 가장 비슷한 임가를 찾아 그들의 "
           "실제 ROI 분포를 보여줍니다. 모델을 신뢰하지 않는 이용자도 “나와 같은 조건인 실제 임가들”은 "
           "납득할 수 있습니다.")
    P(doc, "**기능 7 — 작목 조합 위험 분산.** 금융의 평균–분산 접근을 임산물에 적용했습니다. "
           "상세는 4)절에 기술합니다.")
    P(doc, "**기능 8~11 —** 출하 시기 추천(KAMIS 월별 도매가), 지역 특화도(LQ), 가공 손익분기, "
           "보조사업 지렛대 분석.")

    H2(doc, "다. 핵심 기술")
    TABLE(doc, [
        ["구분", "기술"],
        ["모델", "XGBoost (CUDA, NVIDIA A6000 ×2, tree_method=hist, device=cuda)"],
        ["탐색", f"Optuna TPE multivariate sampler — Model A {A['optuna_xgboost']['n_trials']} trial / "
                f"Model B {B['optuna_xgboost']['n_trials']} trial"],
        ["목적함수", "5-fold CV OOF RMSE (단일 홀드아웃 과적합 방지)"],
        ["불확실성", "Quantile Regression (reg:quantileerror)"],
        ["설명", "TreeSHAP (pred_contribs)"],
        ["최적화", "Dirichlet 무작위 가중 6,000개 + 상관 수축 + 최근접 양정부호 행렬 투영"],
        ["서비스", "FastAPI (Python) + Vue 3 / Vite / Apache ECharts"],
    ], widths=[2.6, 13.8])

    FIGURE(doc, "fig07_pipeline.png", "[그림 3] 분석 파이프라인 전 과정", 12.8)

    # ══════════════════════════════════════════════════════════ 3)
    H1(doc, "3) 데이터 분석 방법")

    H2(doc, "가. 활용 데이터 — 임업통계 (필수)")
    P(doc, "산림청 국가승인통계, 통계청 MDIS 마이크로데이터", 9.6, MUTED, after=4)
    TABLE(doc, [
        ["통계명", "제공기관", "기간", "규모", "용도"],
        ["임가경제조사", "산림청", "2019~2023", f"{A['dataset']['rows']:,} 임가-연도", "Model A 학습"],
        ["임산물생산비조사", "산림청", "2018~2024", f"{B['dataset']['rows']:,} 관측 (5품목)", "Model B 학습"],
        ["임산물생산조사", "산림청", "2022~2024", "196,579 관측", "지역 단가·LQ·가공 분석"],
        ["임업경영실태조사", "산림청", "2018·2020", "임가 단위", "경영 실태 교차 검증"],
    ], widths=[3.6, 2.0, 2.4, 4.2, 4.2])

    H2(doc, "나. 활용 데이터 — 공공데이터 융복합")
    TABLE(doc, [
        ["데이터명", "제공기관", "용도"],
        ["KAMIS 농수산물 도매시장 가격 정보", "한국농수산식품유통공사(aT)", "품목별 월별 도매가 → 출하 시기 추천"],
        ["기상청 API 허브 종관기상관측(ASOS)", "기상청", "시군구 230개 · 기상 지표 15종 결합"],
        ["산림사업 보조금 지원 사업 목록", "산림청", "42개 사업 자부담률 → 실효 ROI 산출"],
        ["행정구역 경계 (시군구)", "국토교통부", "시군구 단위 단가 지도 시각화"],
    ], widths=[6.0, 4.2, 6.2])

    H2(doc, "다. 전처리 상세")
    BUL(doc, "**인코딩·구분자 자동 판별** — 연도별로 CP949/UTF-8, 쉼표/탭이 섞여 있어 자동 감지 로직을 두었습니다.")
    BUL(doc, "**파일설계서(xlsx) 코드북 파싱** — 코드값(1, 2, 3…)을 사람이 읽는 라벨로 되돌리기 위해 설계서를 직접 파싱했습니다.")
    BUL(doc, "**연도별 스키마 정규화** — 조사 연도마다 컬럼명과 코드 체계가 달랐습니다. 특히 "
             "**경영수준별 코드가 2020~2022년은 5/6, 2023~2024년은 1/2로 바뀌어 있었고 이는 "
             "파일설계서에 기재되어 있지 않았습니다.** 분포 대조로 확인해 정합화했습니다.")
    BUL(doc, "**이상치 처리** — 품목별로 IQR 기준을 따로 적용했습니다. 품목마다 수익 분포의 폭이 달라 "
             "전체 일괄 기준을 쓰면 표고류가 통째로 잘려 나갑니다.")
    BUL(doc, "**파생변수 생성** — ha당 경영비·자본·가용노동력, 로그 변환, 경영비/자본 비율, "
             "임업외소득 비중, 지역×업종 교차 범주 등.")

    H2(doc, "라. 데이터 누수 통제 — 본 분석의 방법론적 핵심")
    P(doc, "수익률 예측에서 가장 흔한 오류는 **타깃의 구성요소를 설명변수로 넣는 것**입니다. "
           "ROI = 소득 ÷ 경영비이므로, 소득을 계산하는 데 쓰인 항목이 피처에 남으면 모델은 "
           "예측이 아니라 나눗셈을 학습합니다.")
    P(doc, "**Model A** — 정의상 타깃을 결정하는 7개 변수를 사전 차단했습니다.")
    CODE(doc, '''LEAKY = ["임업총수입", "임업소득", "임가소득", "경상소득",
         "임가순소득", "임가처분가능소득", "임가경제잉여"]''')
    NOTE(doc, "다만 **임업경영비는 남겼습니다.** ROI의 분모이지만, 임가가 연초에 결정하는 "
              "사전(ex-ante) 의사결정 변수이지 결과가 아니기 때문입니다. 이 변수를 빼면 "
              "“경영비를 얼마로 잡을까”라는 가장 실용적인 질문에 답할 수 없게 됩니다.")

    P(doc, "**Model B — 실제로 누수가 발생했고, 이를 자체 적발하여 교정했습니다.**")
    P(doc, "1차 학습 결과가 **Test R² 0.9102**로 나왔습니다. 임업 수익 예측에서 이 수치는 "
           "비현실적이라 판단해 피처–타깃 상관을 전수 감사했습니다. 그 결과 건표고_소계수량과 "
           "ROI의 상관이 **0.885**였습니다. 기존 규칙이 ‘평가액·수확량·생산량’만 잡고 있어 "
           "**총생산액·부가가치·등급별 수량**이 통과하고 있었습니다.")
    P(doc, "정규식 규칙을 확장하여 재학습한 결과, 최고 상관은 **0.885 → 0.344**(노동비 비중)로 "
           "정상화되었고 Test R²는 **0.9102 → 0.6243**으로 내려갔습니다.")
    CODE(doc, '''LEAKY_PATTERNS = [
    r"^소득$", r"^순수익$", r"^부가가치", r"평가액", r"생산액",
    r"수확량", r"생산량", r"수량",
    r"^생산비합계", r"^직접생산비$", r"^간접생산비$", r"^내급비$",
    r"^타용도소비량$", r"^판매량", r"판매금액", r"^총수입",
]
# 투입 수량(비료 시비량 등)은 사전 결정 변수이므로 예외로 남긴다
KEEP_DESPITE_PATTERN = re.compile(r"^(무기질|유기질)_.*_수량_단위당$")''')
    NOTE(doc, "**성능은 0.29 내려갔지만, 그 0.62가 실제로 쓸 수 있는 숫자입니다.** "
              "이 경위를 감추지 않고 보고서와 서비스 화면에 그대로 기재했습니다.", "warn")

    FIGURE(doc, "fig02_leakage.png", "[그림 4] 데이터 누수 자체 적발 및 교정 경위", 12.5)

    H2(doc, "마. 모델링 절차")
    TABLE(doc, [
        ["단계", "내용"],
        ["① 분할", f"층화 8:1:1 (train {A['dataset']['train']:,} / valid {A['dataset']['valid']} / "
                 f"test {A['dataset']['test']}), seed 42 고정"],
        ["② 베이스라인 3종", "동일 분할·동일 지표로 비교 — 집단 평균(지역×업종 71개 집단), 선형회귀, Optuna XGBoost"],
        ["③ 하이퍼파라미터 탐색", "Optuna TPE. 목적함수는 5-fold CV OOF RMSE — 단일 홀드아웃 최적화는 valid에 과적합됨"],
        ["④ GPU 병렬", "NVIDIA A6000 2장에 trial을 라운드로빈 배정"],
        ["⑤ 테스트셋 평가", "최종 파라미터 확정 후 단 한 번만 열었음"],
    ], widths=[3.4, 13.0])

    P(doc, "탐색으로 고른 주요 하이퍼파라미터는 다음과 같습니다. 두 모델의 값이 크게 다른 것은 "
           "자료의 성격 차이를 그대로 반영합니다. 설명변수가 22개뿐인 Model A는 깊게(max_depth 10) "
           "파고들어야 신호를 잡을 수 있고, 169개인 Model B는 얕게 두되 규제를 강하게(reg_lambda 12.1) "
           "걸어야 과적합을 피할 수 있습니다.")
    pa, pb = A["optuna_xgboost"]["best_params"], B["optuna_xgboost"]["best_params"]
    TABLE(doc, [
        ["하이퍼파라미터", "Model A", "Model B", "의미"],
        ["learning_rate", f"{pa['learning_rate']:.4f}", f"{pb['learning_rate']:.4f}", "한 걸음의 보폭"],
        ["max_depth", str(pa["max_depth"]), str(pb["max_depth"]), "나무의 깊이"],
        ["subsample / colsample", f"{pa['subsample']:.2f} / {pa['colsample_bytree']:.2f}",
         f"{pb['subsample']:.2f} / {pb['colsample_bytree']:.2f}", "행·열 표본 비율"],
        ["reg_lambda", f"{pa['reg_lambda']:.3f}", f"{pb['reg_lambda']:.3f}", "L2 규제 강도"],
        ["n_estimators / 탐색", f"{A['optuna_xgboost']['n_estimators']:,} / "
         f"{A['optuna_xgboost']['n_trials']}회",
         f"{B['optuna_xgboost']['n_estimators']:,} / {B['optuna_xgboost']['n_trials']}회",
         "나무 개수 / Optuna 시행"],
    ], widths=[4.0, 3.0, 3.0, 6.4], align_right=[1, 2])

    NOTE(doc, "**분할 감사 — 스스로 의심하고 확인했습니다.** 임가경제조사는 같은 임가를 여러 해 "
              "따라가는 패널입니다(임가 2,168곳 / 행 4,438). 행 단위로 무작위 분할하면 시험셋 "
              "444행 중 **297행(66.9%)이 학습셋과 같은 임가**입니다. 임가번호는 설명변수에 없지만 "
              "지역·업종·규모·경영비·자본의 조합이 사실상 임가를 특정하므로 외워서 맞힐 여지가 "
              "있습니다. 임가 단위로 다시 나눠 재보니 R² 0.1755 → 0.1574로, fold 간 표준편차"
              "(0.0321) 안이었습니다. **영향이 제한적임을 확인했고, 보고한 수치는 그대로 "
              "유효합니다.** 감사 코드는 src/audit_split.py에 있습니다.")

    H2(doc, "바. 활용 프로그램")
    P(doc, "Python 3.11 / XGBoost(CUDA) / Optuna / scikit-learn / pandas / NumPy / SciPy / "
           "FastAPI / Pydantic / Vue 3 / Vite / Apache ECharts / Playwright(화면 검증)")

    # ══════════════════════════════════════════════════════════ 4)
    H1(doc, "4) 분석 내용 및 결과")

    H2(doc, "가. 결과 ① 예측 성능 — 집단 평균 대비 4~6배")
    ab, al, ax_ = (A["forest_service_baseline"]["test"], A["linear_regression"]["test"],
                   A["optuna_xgboost"]["test"])
    bb, bl, bx = (B["forest_service_baseline"]["test"], B["linear_regression"]["test"],
                  B["optuna_xgboost"]["test"])
    TABLE(doc, [
        ["모델", "Model A R²", "Model A RMSE", "Model B R²", "Model B RMSE"],
        ["산림청 방식 (집단 평균)", f"{ab['R2']:.4f}", f"{ab['RMSE']:.2f}", f"{bb['R2']:.4f}", f"{bb['RMSE']:.2f}"],
        ["선형회귀", f"{al['R2']:.4f}", f"{al['RMSE']:.2f}", f"{bl['R2']:.4f}", f"{bl['RMSE']:.2f}"],
        ["**Optuna XGBoost**", f"**{ax_['R2']:.4f}**", f"**{ax_['RMSE']:.2f}**",
         f"**{bx['R2']:.4f}**", f"**{bx['RMSE']:.2f}**"],
        ["개선 배수", "**×4.05**", "−7.1%", "**×5.95**", "−35.2%"],
    ], widths=[4.6, 2.9, 3.0, 2.9, 3.0], align_right=[1, 2, 3, 4])

    FIGURE(doc, "fig01_benchmark.png", "[그림 5] 동일 분할·동일 지표에서의 3종 모델 비교", 12.5)

    NOTE(doc, "**결과의 정직한 해석.** Model A의 R² 0.1736은 절대 수치로 높지 않습니다. 이는 모델의 "
              "한계가 아니라 **임가경제조사 총괄표가 담을 수 있는 정보의 한계**입니다. 산의 수령·수종 "
              "구성·경사·토양·판로 같은 결정적 변수가 총괄표에 없습니다. 실제로 이런 변수를 담은 "
              "임산물생산비조사(Model B)에서는 R²가 0.62까지 올라갑니다. "
              "**두 모델의 격차 자체가 “어떤 자료를 더 모아야 하는가”에 대한 답입니다.**", "warn")

    P(doc, "Model B를 품목별로 나누어 보면 성능이 표본 수를 그대로 따라갑니다. "
           "**표고 노지는 R²가 음수(−0.052)입니다.** 테스트 표본이 32건뿐이라 모델이 평균보다도 "
           "못 맞힌다는 뜻입니다. 감추지 않고 그대로 싣고, 서비스 화면에서도 해당 품목에는 "
           "“표본이 적어 참고만 하십시오”라고 표시합니다.")
    pit = B["per_item_test"]
    rows = [["품목", "Test R²", "RMSE", "MAE", "테스트 표본", "판정"]]
    for k, v in sorted(pit.items(), key=lambda kv: -kv[1]["R2"]):
        judge = "양호" if v["R2"] > 0.3 else ("제한적" if v["R2"] > 0 else "표본 부족")
        rows.append([k, f"{v['R2']:.4f}", f"{v['RMSE']:.1f}", f"{v['MAE']:.1f}",
                     f"{v['n']}건", judge])
    TABLE(doc, rows, widths=[2.6, 2.4, 2.4, 2.4, 2.6, 2.8], align_right=[1, 2, 3, 4])

    H2(doc, "나. 결과 ② 패널 구조를 쓰면 설명력이 네 배가 된다")
    P(doc, "임가경제조사가 같은 임가를 여러 해 따라가는 패널이라는 점을 확인하고, "
           "**직전 연도 실적을 설명변수로 넣어** 다시 학습했습니다. 임가가 올해를 계획할 때 "
           "작년 결과는 이미 알고 있는 정보이므로 미래 정보를 끌어다 쓰는 것이 아닙니다.")
    pn = json.load(open(os.path.join(ROOT, "models", "metrics_panel.json"), encoding="utf-8"))
    gp, gsd = pn["임가단위_5회평균"], pn["임가단위_fold표준편차"]
    rows = [["모델", "Test R²", "RMSE", "MAE", "fold 편차"]]
    for k in ("산림청 방식", "기존 22변수", "작년값 선형보정", "패널 변수 추가"):
        lab = f"**{k}**" if k == "패널 변수 추가" else k
        rows.append([lab, f"{gp[k]['R2']:.4f}", f"{gp[k]['RMSE']:.2f}",
                     f"{gp[k]['MAE']:.2f}", f"±{gsd[k]:.4f}"])
    TABLE(doc, rows, widths=[4.4, 3.0, 2.8, 2.8, 2.8], align_right=[1, 2, 3, 4])
    P(doc, "같은 부분집합에서 다시 잰 값입니다. **작년 ROI 한 항목을 선형 보정한 것만으로 "
           "R² 0.2548**이 나와, 설명변수 22개를 쓴 기존 모델(0.0635)의 네 배입니다. "
           "패널 변수를 넣은 XGBoost는 0.2798로 **산림청 방식(0.0274) 대비 10.2배**입니다.")
    NOTE(doc, "**개선의 대부분은 모델이 아니라 자료에서 나왔습니다.** 작년값 선형보정만으로 "
              "0.2548이고, 모델이 그 위에 더한 몫은 +0.025에 불과합니다. 이 점을 감추지 "
              "않습니다. 여기서 읽어야 할 것은 알고리즘의 우수성이 아니라 "
              "**임가 수익률에는 조사표에 잡히지 않는 고유한 몫이 크고, 작년 실적이 그 대리 "
              "지표 노릇을 한다**는 사실입니다.")
    P(doc, "과거로 배워 미래를 맞히는 조건 — 실제 사용 상황과 같습니다 — 에서도 앞섭니다. "
           "2020년까지 배워 2021년을 맞힐 때 R² 0.2200(기존 0.1404), 2021년까지 배워 "
           "2022년을 맞힐 때 0.3207(기존 0.2442)입니다.")
    FIGURE(doc, "fig14_panel.png",
           "[그림 6] 패널 구조 활용 — 임가 단위 분할(좌)과 연도 단위 분할(우)", 14.9)
    NOTE(doc, "**적용 범위와 한계.** 이 모델은 직전 연도 관측이 있는 행에만 쓸 수 있습니다"
              "(4,438행 중 2,186행, 49.3%). 임가번호가 2019~2022년은 네 자리(1001…), "
              "2023년은 다섯 자리(11021…)로 바뀌어 **2022년과 2023년 사이 연결이 완전히 "
              "끊깁니다(교집합 0곳).** 파일설계서에 안내가 없어 대조로 확인했습니다. "
              "대조표가 마련되면 2023년 이후에도 쓸 수 있고, 그것만으로 예측 정확도가 "
              "네 배 오릅니다. 처음 이용하는 임가에게는 기존 Model A를 씁니다.", "warn")

    H2(doc, "다. 결과 ③ 예측구간이 실제로 맞는다")
    TABLE(doc, [
        ["모델", "구간 포함률 (명목 80%)", "P10 미만", "P90 초과", "구간 폭 중앙값"],
        ["Model A", f"**{Q['roi']['coverage_80pct']*100:.1f}%**",
         f"{Q['roi']['coverage_below_p10']*100:.1f}%", f"{Q['roi']['coverage_above_p90']*100:.1f}%",
         f"{Q['roi']['median_interval_width']:.1f}%p"],
        ["Model B", f"**{Q['cost']['coverage_80pct']*100:.1f}%**",
         f"{Q['cost']['coverage_below_p10']*100:.1f}%", f"{Q['cost']['coverage_above_p90']*100:.1f}%",
         f"{Q['cost']['median_interval_width']:.1f}%p"],
    ], widths=[2.8, 4.4, 2.8, 2.8, 3.6], align_right=[1, 2, 3, 4])
    P(doc, "양쪽 꼬리가 각각 10~13%로 고르게 분포합니다. 구간이 한쪽으로 치우치지 않았다는 뜻입니다.")


    H2(doc, "라. 결과 ④ 사례 연구 — 한 임가에게 실제로 무엇을 말해 주는가")
    P(doc, "지표만 늘어놓으면 이 분석이 현장에서 어떻게 쓰이는지 보이지 않습니다. "
           "실제 입력에 대한 시스템의 출력을 그대로 옮깁니다.")
    TABLE(doc, [
        ["입력 항목", "값"],
        ["지역 / 업종", "충남 / 밤재배업"],
        ["임지 규모 / 전겸업", "5~10ha 미만 / 임업주업"],
        ["경영주 연령 / 가구원", "60대 / 3명"],
        ["연간 임업경영비", "1,500만원"],
        ["기초 자본(순재산) / 연초 보유", "3억원 / 500만원"],
    ], widths=[5.6, 10.8])
    P(doc, "**① 예측** — ROI **111.0%** (P10 −48.9% ~ P90 286.9%). "
           "집단 평균 방식은 같은 조건에 233.2%를 제시하지만, 이 임가의 경영비·자본 구조를 "
           "반영하면 그보다 낮게 나옵니다. 구간이 넓다는 사실 자체가 “확정된 수익이 아니다”라는 "
           "정보입니다.")
    P(doc, "**② 설명** — 전체 평균 152.7%에서 이 임가의 111.0%까지, 무엇이 얼마나 끌어내렸는지 "
           "가법 분해합니다. 경영비 −63.5%p, 임지 규모 −41.5%p가 주된 하락 요인이고, "
           "작목(+31.1%p)과 전업 여부(+22.1%p)가 이를 일부 상쇄합니다.")
    P(doc, "**③ 처방** — 같은 조건 임가의 중앙값(152.3%)을 목표로 반사실 탐색을 수행한 결과, "
           "**한 해 경영비를 1,500만원에서 750만원으로 조정하면 예측 ROI는 195.9%**로 오릅니다. "
           "다만 이는 비율의 개선이며 소득 총액과는 다른 문제라는 점을 화면에 함께 표시합니다.")
    FIGURE(doc, "fig12_case.png",
           "[그림 7] 사례 — 예측 근거의 분해(좌)와 반사실 처방(우)", 12.8)
    P(doc, "**④ 비교** — 조건이 가장 비슷한 임가 40곳의 ROI 중앙값은 146.2%이며, 그중 잘 버는 "
           "쪽은 247.6%입니다. 두 집단을 가르는 가장 큰 차이는 **면적당 일손(+300%)**과 "
           "**면적당 쓰는 돈(+270.9%)**이었습니다. 같은 금액을 쓰더라도 더 좁은 면적에 집중해 "
           "투입하는 쪽이 수익률이 높다는 뜻으로, 앞의 처방과 방향이 일치합니다.")
    NOTE(doc, "이 네 단계가 본 시스템의 핵심입니다. **예측만 주면 임가는 믿을 이유가 없고, "
              "설명만 주면 무엇을 할지 모릅니다.** 예측·설명·처방·비교를 함께 제시해야 "
              "비로소 의사결정에 쓸 수 있습니다.")

    H2(doc, "마. 결과 ⑤ 품질 관리의 수익 기여를 금액으로 환산")
    g, sim = IN["등급별_단가"]["밤"], IN["등급전환_시뮬레이션"]["밤"]
    P(doc, f"밤의 등급 간 단가 격차는 **최고/최저 {g['최고_최저_단가배수']:.2f}배**입니다. "
           f"물량 보존(같은 산물의 등급 간 이동)을 가정하고 **{sim['전환_시나리오']}** 시 "
           f"ha당 수취액은 **+{sim['수취액_증가_원per단위면적']:,.0f}원** 증가합니다. "
           f"단가차 {sim['단가차_원']:,.0f}원 × 전환 물량으로 계산한 값입니다.")
    age = IN["수령별_수익성"]
    P(doc, f"수령별로는 밤 {age['밤']['최고구간']}(ROI {age['밤']['최고ROI']:.0f}%)이 정점이며 "
           f"**{age['밤']['최저구간']} 구간에서 ROI {age['밤']['최저ROI']:.0f}%로 역전**됩니다. "
           f"떫은감은 {age['떫은감']['최고구간']}({age['떫은감']['최고ROI']:.0f}%), "
           f"대추는 {age['대추']['최고구간']}({age['대추']['최고ROI']:.0f}%)이 정점입니다. "
           "임목 갱신 판단의 정량적 근거가 됩니다.")

    FIGURE(doc, "fig06_insight.png", "[그림 8] 등급 전환 효과 및 수령 구간별 수익성", 12.5)

    H2(doc, "바. 결과 ⑥ 지역 — 어디서 무엇이 잘 되고, 어디가 값을 더 받는가")
    P(doc, "임산물생산조사 196,579개 관측을 시군구 단위로 집계해 두 가지를 봤습니다. "
           "하나는 **어느 지역에서 어떤 품목이 실제로 수익을 내는가**이고, 다른 하나는 "
           "**같은 품목이라도 지역에 따라 단가가 얼마나 다른가**입니다.")
    FIGURE(doc, "fig11_heatmap.png",
           "[그림 9] 지역×품목 수익성 — 조사 임가 15곳 이상인 조합만 표시", 11.4)
    P(doc, "빈 칸이 많다는 사실 자체가 정보입니다. 대부분의 지역·품목 조합은 표본이 15곳에 "
           "미치지 못해 판단할 수 없습니다. 그런 칸에 억지로 숫자를 채우지 않았습니다.")
    P(doc, "특화도는 입지계수(LQ)로 쟀습니다. 그 지역 산업 구성에서 해당 품목이 전국 평균보다 "
           "얼마나 큰 비중을 차지하는지 보는 지표로, 1을 넘으면 특화되었다고 봅니다. "
           "**밤의 충남 LQ는 6.73**으로 전국 생산금액의 59.1%가 한 시도에 몰려 있습니다"
           "(세종 3.23, 충북 1.32). 단가는 별개입니다. 2024년 밤의 전국 가중평균단가 "
           "2,699원/kg 대비 경기가 +30.4%, 대전 +28.9%, 충북 +27.2%로 **주산지가 반드시 "
           "단가가 높지는 않습니다.** 물량이 몰리는 곳은 오히려 단가가 눌립니다. "
           "생산 특화도와 단가 프리미엄을 함께 봐야 출하처 판단이 가능합니다.")

    H2(doc, "사. 결과 ⑦ 작목 조합 위험 분산, 그리고 예상 밖의 발견")
    P(doc, "임업 조언은 대개 “무엇이 가장 돈이 되는가”만 답합니다. 그러나 임가가 실제로 겪는 문제는 "
           "**그해 시세와 작황이 흔들린다**는 것입니다. 금융의 평균–분산 접근을 임산물 ROI에 "
           "적용했습니다.")
    P(doc, "**위험을 어떻게 쟀는가 — 이 대목이 결정적입니다.** ROI가 흩어지는 원인은 두 가지입니다.")
    BUL(doc, "**(가) 연도 효과** — 그해 시세·작황. 모든 임가가 함께 겪습니다.")
    BUL(doc, "**(나) 임가 효과** — 같은 해, 같은 작목 안에서도 임가마다 다릅니다.")
    P(doc, "**작목을 섞어 줄일 수 있는 것은 (가)뿐입니다.** (나)는 그 임가의 특성이라 어느 작목을 "
           "하든 따라옵니다. 밤과 대추를 반씩 해도 서툰 사람은 양쪽 다 서툽니다.")
    NOTE(doc, "초기 계산에서 임가 간 산포(σ 117~196%p)를 그대로 위험으로 놓았더니 조합 변동폭이 "
              "**3.4%p**라는 비현실적 값이 나왔습니다. √n 효과의 허구입니다. "
              "연도 효과(σ 10.7~39.5%p)만 분산 대상으로 재정의하여 교정했습니다.", "warn")
    P(doc, "또한 표고류는 2018~2022년, 과실류는 2020~2024년 자료라 **겹치는 해가 3년뿐**입니다. "
           "점 3개로 잰 상관은 우연히 ±1에 가깝게 나오기 쉬워(실제 대추–표고 노지 −0.98), 겹친 해 "
           "수에 따라 0쪽으로 수축(λ = (n−1)/(n−1+8))시키고, 그래도 행렬이 성립하지 않는 경우 "
           "**최근접 양정부호 행렬로 투영**했습니다.")

    rows = [["작목", "기대수익", "연도 변동 σ", "임가 격차 σ", "효율", "표본", "조사연도"]]
    for r in sorted(PF["품목"], key=lambda x: -(x["효율"] or 0)):
        rows.append([r["품목"], f"{r['기대수익_pct']:.1f}%", f"±{r['연도변동_pct']:.1f}%p",
                     f"±{r['임가격차_pct']:.1f}%p", f"{r['효율']:.2f}",
                     f"{r['표본수']:,}", f"{r['조사연도'][0]}~{r['조사연도'][-1]}"])
    TABLE(doc, rows, widths=[2.4, 2.4, 2.6, 2.6, 1.8, 2.0, 2.6], align_right=[1, 2, 3, 4, 5])

    bp = PF["최고효율"]
    P(doc, "**최적 조합** — " + " · ".join(f"{c['품목']} {c['비중_pct']}%" for c in bp["구성"])
           + f" → 기대수익 {bp['기대수익_pct']}%, 연도 변동 **±{bp['연도변동_pct']}%p**, "
             f"효율 {bp['효율']}. 단일 최고효율({PF['단일_최고효율']['품목']} "
             f"{PF['단일_최고효율']['효율']}) 대비 **{PF['분산효과_pct']:+.1f}%** 개선. "
             f"작목 간 상관을 일률적으로 0.3으로 보는 보수적 가정에서도 효율 "
             f"{PF['보수가정']['효율']}로 결론이 유지됩니다.")

    FIGURE(doc, "fig04_portfolio.png", "[그림 10] 작목 조합의 위험–수익 지도 (평균–분산 접근)", 10.3)

    ratio = bp["임가격차_pct"] / bp["연도변동_pct"]
    NOTE(doc, f"**그런데 더 중요한 발견이 있습니다.** 최적 조합의 연도 변동은 ±{bp['연도변동_pct']}%p인데, "
              f"**임가 격차는 ±{bp['임가격차_pct']}%p로 {ratio:.1f}배**입니다. "
              "**무엇을 심을지 고르는 것보다 어떻게 하느냐가 수익을 훨씬 크게 가릅니다.** "
              "이것은 본 과제의 출발점 — “집단 평균이 아니라 임가별로 계산해야 한다” — 를 데이터가 "
              "스스로 증명한 결과입니다. 부수적 발견이 오히려 프로젝트 전체의 논지를 뒷받침하게 "
              "되었습니다.")
    lead = IN["선도임가_격차"]["밤"]
    P(doc, f"실제로 밤 기준 선도임가({lead['표본']['선도임가']}곳)의 ROI 중앙값은 "
           f"**{lead['선도임가']['ROI중앙값']}%**, 그 외({lead['표본']['이외임가']}곳)는 "
           f"**{lead['이외임가']['ROI중앙값']}%**입니다. 같은 품목, 같은 해에도 이만큼 벌어집니다.")

    FIGURE(doc, "fig05_risk_split.png", "[그림 11] 수익 분산의 원천 분해 — 임가 효과 대 연도 효과", 11.3)

    H2(doc, "아. 결과 ⑧ 그 밖의 분석")
    BUL(doc, "**지역 특화도(LQ)** — 밤은 충남 LQ 6.73(전국 생산금액의 59.1%)으로 압도적 주산지.")
    BUL(doc, "**가공 손익분기** — 생표고 8.0kg가 건표고 1kg이 되므로 단가 배수가 8.0배를 넘어야 "
             "가공이 유리한데, **실제 배수는 4.13배**. 원물 직판 대비 −48.4%로 **원물 직판이 유리**. "
             "“가공하면 부가가치가 오른다”는 통념을 자료로 반박한 사례입니다.")
    BUL(doc, "**보조사업 지렛대** — 42개 사업의 자부담률로 실효 ROI 산출. 수출 분야는 자부담 0%"
             "(최대 지렛대 10배), 생산기반은 최저 20%(5배).")
    BUL(doc, "**기상 결합** — 시군구 230개 기상 지표 15종을 결합했으나 예측력 개선은 미미했습니다. "
             "**실패한 시도도 그대로 기록했습니다.** 임가경제조사가 시도(9개) 단위여서 기상의 국지성이 "
             "희석되기 때문으로 판단하며, 이는 **조사 설계에 대한 시사점**입니다.")

    H2(doc, "자. 결과 ⑨ 서비스 구현")
    P(doc, "FastAPI + Vue 3 웹 애플리케이션으로 전 기능을 제공합니다. 임가가 조건을 바꾸면 모든 화면이 "
           "즉시 재계산됩니다. 화면 구성은 **임업 종사자가 쓸 수 있는 말**로 전면 재작성했습니다.")
    TABLE(doc, [
        ["통계 용어", "화면 표기"],
        ["ROI 예측", "얼마나 남을까"],
        ["출하 시기 최적화", "언제 팔까"],
        ["반사실 처방", "무엇을 바꾸면 좋을까"],
        ["포트폴리오 최적화", "한 작목에 몰아도 될까"],
    ], widths=[6.0, 6.0])
    P(doc, "모든 수치에 **출처·표본수·한계**를 병기했습니다.")

    # ══════════════════════════════════════════════════════════ 5)
    H1(doc, "5) 주요 성과 및 기대효과")

    H2(doc, "가. 데이터 분석 활용 전 / 후")
    TABLE(doc, [
        ["데이터 분석 활용 전", "⇨", "데이터 분석 활용 후"],
        ["지역·품목 집단 평균 제공 (R² 0.0428)", "⇨", "임가별 개인화 예측 (R² 0.1736 / 0.6243, **4.05~5.95배**)"],
        ["단일 숫자 제시, 불확실성 은폐", "⇨", "P10~P90 예측구간, **포함률 78.2% / 76.5% 검증**"],
        ["“평균이 이렇다”는 현황 기술", "⇨", "“무엇을 얼마로 바꾸면 얼마가 된다”는 **반사실 처방**"],
        ["예측 근거 불투명", "⇨", "**TreeSHAP 변수별 기여 분해**"],
        ["단일 작목 수익만 비교", "⇨", "**위험 대비 수익 기준 작목 조합** (효율 +33.3%)"],
        ["마이크로데이터가 연구자 전유물", "⇨", "**웹에서 누구나 즉시 이용**"],
        ["통계 용어 중심 보고서", "⇨", "**임업 종사자의 말로 재작성**"],
    ], widths=[6.4, 1.0, 9.0])

    H2(doc, "나. 정책적 시사점")
    P(doc, "**① 조사 설계 개선 근거.** Model A(R² 0.17)와 Model B(R² 0.62)의 격차는 "
           "**임가경제조사 총괄표에 수령·수종·경사·판로 정보가 없어서** 생깁니다. 이 항목들을 보강하면 "
           "개인화 예측 정확도가 크게 오릅니다. 본 분석은 그 필요성을 **수치로 입증한 사례**입니다.")
    P(doc, f"**② 지원 정책의 초점 이동.** 임가 격차(±{bp['임가격차_pct']}%p)가 작목 선택 위험"
           f"(±{bp['연도변동_pct']}%p)의 **{ratio:.1f}배**입니다. 품목 전환 지원보다 "
           "**경영 역량 격차 해소**(선도임가 기술 이전, 품질 관리 교육)가 효율이 높다는 것을 "
           f"시사합니다. 실제로 선도임가와 일반 임가의 ROI는 {lead['선도임가']['ROI중앙값']}% 대 "
           f"{lead['이외임가']['ROI중앙값']}%로 벌어져 있습니다.")
    P(doc, f"**③ 품질 관리의 정량적 근거.** 등급 10%p 전환만으로 ha당 "
           f"+{sim['수취액_증가_원per단위면적']:,.0f}원. 선별·전정 지원 사업의 비용편익 산정에 "
           "바로 쓸 수 있습니다.")
    P(doc, "**④ 패널 연결의 복구 — 비용 대비 효과가 가장 큰 개선.** 임가번호 체계가 2023년에 "
           "바뀌면서 패널 연결이 끊겼습니다. **연도 간 대조표를 마련하는 것만으로 예측 "
           "설명력이 0.06에서 0.28로, 네 배 이상 오릅니다.** 새 조사 항목을 추가하는 것보다 "
           "훨씬 적은 비용으로 얻을 수 있는 개선입니다. 임가경제조사가 이미 패널로 설계되어 "
           "있는데도 그 구조가 활용되지 못하고 있는 셈입니다.")
    P(doc, "**⑤ 가공 지원의 재검토.** 표고 가공은 손익분기 8.0배에 실제 4.13배로 미달합니다. "
           "가공 시설 지원이 항상 유리하지 않다는 것을 보여줍니다.")

    H2(doc, "다. 현장 활용 시나리오")
    TABLE(doc, [
        ["시점", "임가가 하는 일", "시스템이 답하는 것"],
        ["연초 · 영농 계획", "올해 경영비를 얼마로 잡을지 정한다",
         "경영비 구간별 예측 ROI 곡선과 같은 조건 임가의 중앙값"],
        ["봄 · 투자 판단", "임목 갱신·시설 도입을 검토한다",
         "수령 구간별 수익성, 보조사업 자부담률 기반 실효 ROI"],
        ["여름 · 품질 관리", "선별·전정에 얼마나 품을 들일지 정한다",
         "등급 전환 시 ha당 수취액 증가분(밤 기준 +75,178원)"],
        ["가을 · 출하", "언제, 어디로 낼지 정한다",
         "KAMIS 월별 도매가 추이, 시군구 단가 프리미엄"],
        ["연말 · 복기", "올해 결과가 어땠는지 확인한다",
         "유사 임가 40곳과의 비교, 잘 버는 쪽과의 투입 구조 차이"],
    ], widths=[2.8, 5.6, 8.0])
    P(doc, "임업은 결정과 결과 사이의 간격이 긴 산업입니다. 밤나무를 새로 심으면 수확까지 여러 해가 "
           "걸립니다. 그래서 **결정 시점에 근거를 주는 것**이 사후 통계를 정확히 만드는 것보다 "
           "임가에게 실질적인 도움이 됩니다.")

    H2(doc, "라. 기대효과")
    TABLE(doc, [
        ["대상", "기대효과"],
        ["임가", "투자 전 수익 구간 확인 → 과잉 투자 방지 / 임목 갱신·출하 시기 판단 근거 확보"],
        ["산림청·진흥원", "조사 항목 개선 우선순위 도출 / 지원 사업 비용편익 산정 근거"],
        ["지자체", "지역 특화도(LQ) 기반 품목 육성 전략 / 시군구 단위 단가 비교"],
        ["연구자", "전 과정 재현 가능한 코드 공개 → 후속 연구 기반"],
    ], widths=[3.0, 13.4])

    H2(doc, "마. 확장 계획")
    BUL(doc, "미확보 통계 3종(임산물소득조사·산림산업조사·산림휴양복지활동조사) 추가 결합")
    BUL(doc, "임가경제조사 세부 파일 결합으로 Model A 설명력 보강")
    BUL(doc, "산림경영지도원 현장 상담 도구로 배포")

    H2(doc, "바. 재현성 및 별첨")
    P(doc, "본 과제의 모든 결과는 코드로 재현됩니다. seed는 42로 고정했고, 전처리부터 학습·평가·"
           "산출물 생성까지 스크립트로 구성했습니다. **요강에 따라 소스코드 전체를 분량 외 별첨으로 "
           "제출합니다.**")
    TABLE(doc, [
        ["구분", "파일"],
        ["전처리", "src/preprocess.py, src/preprocess_cost.py"],
        ["학습", "src/train_optuna.py, src/train_cost.py, src/train_quantile.py, src/train_panel.py"],
        ["분석", "src/insights.py, src/production.py, src/management.py, src/subsidy.py, "
               "src/region_map.py, src/portfolio.py, src/explain.py"],
        ["외부 API", "src/kamis_client.py, src/kma_client.py, src/weather.py, src/weather_sgg.py"],
        ["검증", "src/audit_split.py — 분할 감사"],
        ["도표 생성", "src/make_figures.py, src/make_docx.py"],
        ["서비스", "api/ (FastAPI), web/ (Vue 3)"],
    ], widths=[2.4, 14.0])

    H2(doc, "사. 이 분석에서 지킨 원칙")
    NOTE(doc, "① **성능보다 정직을 택했습니다.** Model B의 R²를 0.91에서 0.62로 내린 것이 "
              "이 프로젝트에서 가장 중요한 결정이었습니다.\n"
              "② **한계를 감추지 않았습니다.** 기상 결합 실패, Model A의 낮은 설명력, "
              "표고류 자료의 짧은 기간을 모두 본문에 적었습니다.\n"
              "③ **쓸 수 없는 조언은 하지 않았습니다.** “산을 줄이면 수익률이 오른다”는 결과를 "
              "통계적으로 유의하다는 이유만으로 제시하지 않았습니다.")

    doc.save(OUT)
    print(f"[saved] {OUT}")
    print(f"  문단 {len(doc.paragraphs)}개 · 표 {len(doc.tables)}개")


if __name__ == "__main__":
    build()
